from flask import Flask, request, jsonify
from clarifai_grpc.channel.clarifai_channel import ClarifaiChannel
from clarifai_grpc.grpc.api import service_pb2_grpc
from langchaincoexpert.llms import Clarifai
from langchaincoexpert.agents import load_tools
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import PyPDF2
import re
# import csv
import spacy
import docx2txt
from pprint import pprint
from fpdf import FPDF
from langchaincoexpert.agents import initialize_agent
from langchaincoexpert.utilities import GoogleSearchAPIWrapper# import csv
from langchaincoexpert.agents import AgentType
import fitz

from dropbox_sign import \
    ApiClient, ApiException, Configuration, apis, models
import subprocess

from langchaincoexpert.memory import VectorStoreRetrieverMemory
from langchaincoexpert.chains import ConversationChain
from langchaincoexpert.prompts import PromptTemplate
from langchaincoexpert.vectorstores import SupabaseVectorStore
from langchaincoexpert.embeddings import ClarifaiEmbeddings
from supabase.client import  create_client
# from dotenv import load_dotenv
# from firestore import db
from firebase_admin import firestore
# from dotenv import load_dotenv
import os
from flask_cors import CORS

# Load environment variables from .env file
# load_dotenv()
app = Flask(__name__)

CORS(app, supports_credentials=True)

# Clarifai settings
CLARIFAI_PAT = os.getenv("CLARIFAI_PAT")

# Supabase settings
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_KEY")
supabase = create_client(SUPABASE_URL, SUPABASE_KEY)

#SERPER_API_KEY = os.getenv("SERPER_API_KEY")

os.environ['SERPER_API_KEY']
tools = load_tools(["google-serper"])


nlp = spacy.load("en_core_web_sm")

# Set up the Clarifai channel
channel = ClarifaiChannel.get_grpc_channel()
stub = service_pb2_grpc.V2Stub(channel)

# Clarifai settings
USER_ID = 'ahmedz'
APP_ID = 'FINGU'
MODEL_ID = 'GPT-3_5-turbo'

#Drop Box Config
# configuration = Configuration(
#     # Configure HTTP basic authorization: api_key
#     username="e88383d78903e3ee97788a3993bad96903e846c64052648059096b50b1017f15",

#     # or, configure Bearer (JWT) authorization: oauth2
#     # access_token="YOUR_ACCESS_TOKEN",
# )


# Initialize Clarifai embeddings
embeddings = ClarifaiEmbeddings(pat=CLARIFAI_PAT, user_id="openai", app_id="embed", model_id="text-embedding-ada")

# Initialize Supabase vector store
# vectordb = SupabaseVectorStore.from_documents({}, embeddings, client=supabase)

# Initialize Clarifai LLM
llm = Clarifai(pat=CLARIFAI_PAT, user_id='meta', app_id='Llama-2', model_id='llama2-7b-chat')


# Handle incoming messages
def handle_message(input_text , user_id,internet,spell,assesment):
    memory_key = {user_id}
    if internet : 

        response =  generate_Internet_response_llmchain(input_text, user_id)
    else: 
        response = generate_response_llmchain(input_text, user_id,spell=spell,assessment=assesment)

    
    return response


def generate_Internet_response_llmchain(prompt, conv_id):
    convid = "a" + str(conv_id)
    vectordb = SupabaseVectorStore.from_documents({}, embeddings, client=supabase,user_id=conv_id) # here we use normal userid "for saving memory"

    retriever = vectordb.as_retriever(search_kwargs=dict(k=15,user_id=convid)) # here we use userid with "a" for retreiving memory
    memory= VectorStoreRetrieverMemory(retriever=retriever , memory_key="chat_history")
    DEFAULT_TEMPLATE = """The following is a friendly conversation between a human and an AI called ContractGPT. 
   ,The Ai is a Contract Creation assitant designed to make Contracts.
   If the AI does not know the answer to a question, it truthfully says it does not know or reply with the same question.
    The AI should act as a tool that only outputs a contract results without explanations or comments, and only asks questions when needed too, and always return the whole contract/agreement not parts of it.
{history}
(You do not need to use these pieces of information if not relevant)

Current conversation:
Human: {input}
AI:"""

  
    agent = initialize_agent(tools, llm, agent=AgentType.CONVERSATIONAL_REACT_DESCRIPTION, verbose=True, memory = memory)
    # agent.input_keys= 
    final = agent.run(input = prompt)
    return final



def generate_response_llmchain(prompt, conv_id,spell,assessment):

    convid = "a" + str(conv_id)
    # filter = {"user_id": userid}
    vectordb = SupabaseVectorStore.from_documents({}, embeddings, client=supabase,user_id=conv_id) # here we use normal userid "for saving memory"

    retriever = vectordb.as_retriever(search_kwargs=dict(k=10,user_id=convid)) # here we use userid with "a" for retreiving memory
    memory = VectorStoreRetrieverMemory(retriever=retriever, memory_key=convid)
    if spell:
     DEFAULT_TEMPLATE = """
Relevant pieces of previous conversation:
{user_id}

The Ai's Role is only to fix spelling and grammatical mistakes regardless of anything , it should return the same human input with spelling and grammar fixed.
Current conversation:
Human: {input}
AI:"""

    elif assessment:
        DEFAULT_TEMPLATE = """The following is a friendly conversation between a human and an AI called ContractGPT. 
    The AI Should only make an overall risk assessment to the contract and give notes and advices.
Relevant pieces of previous conversation:
{user_id}
(You do not need to use these pieces of information if not relevant)

Current conversation:
Human: {input}
AI:"""
    else:
        DEFAULT_TEMPLATE = """<s>[INST]The following is a friendly conversation between a human and an AI called ContractGPT. 
   ,The Ai is a Contract Creation assistant designed to make Contracts.
   If the AI does not know the answer to a question, it truthfully says it does not know or reply with the same question.
   The AI should act as a tool that only outputs a contract results without explanations or comments, and only asks questions when needed too, and always return the whole contract/agreement not parts of it.
<<SYS>>
Relevant pieces of previous conversation:
{user_id}
(You do not need to use these pieces of information if not relevant)

Current conversation:
Human: {input}
AI:<</SYS>>[/INST]"""
    
    
 
    formatted_template = DEFAULT_TEMPLATE.format(user_id="{"+convid+"}",input = "{input}")
 
    PROMPT = PromptTemplate(
        input_variables=[convid, "input"], template=formatted_template
    )

    
    conversation_with_summary = ConversationChain(
        llm=llm,
        prompt=PROMPT,
        memory=memory,
        verbose=True
    )

    ans = conversation_with_summary.predict(input=prompt)
    # response = ans
    return ans





def doc_to_text(doc_file_path):
    # Extract text from the .doc file
    text = docx2txt.process(doc_file_path)
    
    return text

def pdf_to_text(pdf_file_path):
    text = ""

    with open(pdf_file_path, 'rb') as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)  # Use PdfReader instead of PdfFileReader

        for page in pdf_reader.pages:
            text += page.extract_text()

    return text

def text_to_docx(text):
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Define a regular expression pattern to find **word** occurrences
    pattern = r'\*\*(.*?)\*\*'

    start = 0

    for match in re.finditer(pattern, text):
        end = match.start()
        # Add text before the **word**
        run = paragraph.add_run(text[start:end])

        start, end = match.span()
        # Add the bold and capitalized word
        word = match.group(1).upper()  # Capitalize the word
        run = paragraph.add_run(word)
        run.bold = True
        run.font.size = Pt(12)  # Set the font size
        start = end

    # Add any remaining text (after the last **word**)
    run = paragraph.add_run(text[start:])

    docx_file_path = "output.docx"
    doc.save(docx_file_path)
    return docx_file_path


def docx_to_text(docx_file_path):
    doc = Document(docx_file_path)
    text = ""

    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"

    return text

# input_text = "This is a **test** sentence. It should **work** fine."
# output_file = text_to_docx(input_text)
    
# print(f"Word document generated at: {output_file}")
# def text_to_pdf(text):
#     pdf = FPDF()
#     pdf.add_page()
#     pdf.set_font("Arial", size=12)
#     pdf.multi_cell(190, 10, txt=text, align="L")
#     pdf_file_path = "output.pdf"
#     pdf.output(pdf_file_path)
#     return pdf_file_path

@app.route('/delete', methods=['DELETE'])
def deleteChat():
    data = request.get_json()
    conversationId = "a" + data['conversationId']
    
    # Check if conversationId is provided
    if not conversationId:
        return jsonify({"message": "Invalid conversation ID provided"}, status_code=400)

    # Initialize the SupabaseVectorStore instance
    vectordb = SupabaseVectorStore.from_documents({}, embeddings, client=supabase, user_id="")

    # Delete the chat based on conversationId
    try:
        vectordb.delete(ids=[conversationId])
        return jsonify({"message": "Chat deleted successfully"})
    except Exception as e:
        return jsonify({"message": f"Error deleting chat: {str(e)}"}, status_code=500)




# @app.route('/convert', methods=['POST'])
# def convert():
#     try:
#         # Get the uploaded file from the request
#         uploaded_file = request.files['file']

#         # Check if a file was provided in the request
#         if not uploaded_file:
#             return jsonify({"error": "No file provided in the request"}, status_code=400)

#         # Save the uploaded DOCX file temporarily
#         temp_docx_path = "temp.docx"
#         uploaded_file.save(temp_docx_path)

#         # Use the docx_to_text function to extract text from the DOCX file
#         extracted_text = docx_to_text(temp_docx_path)

#         # Remove the temporary file
#         os.remove(temp_docx_path)

#         return jsonify({"text": extracted_text})
#     except Exception as e:
#         return jsonify({"error": str(e)}, status_code=500)
def is_docx(filename):
    return filename.lower().endswith('.docx')
def is_doc(filename):
    return filename.lower().endswith('.doc')
def is_pdf(file_extension):
    return file_extension.lower() == '.pdf'

@app.route('/convert', methods=['POST'])
def convert():
    try:
        # Get the uploaded file from the request
        uploaded_file = request.files['file']

        # Check if a file was provided in the request
        if not uploaded_file:
            return jsonify({"error": "No file provided in the request"}, status_code=400)

        # Get the filename and file extension
        filename, file_extension = os.path.splitext(uploaded_file.filename)

        # Save the uploaded file temporarily
        temp_file_path = "temp_file" + file_extension
        uploaded_file.save(temp_file_path)

        # Check if it's a DOCX file based on the file extension
        if is_docx(file_extension):
            extracted_text = docx_to_text(temp_file_path)
        elif is_doc(file_extension):
            extracted_text = doc_to_text(temp_file_path)    
        elif is_pdf(file_extension):

            # Assume it's a PDF file based on the file extension
            extracted_text = pdf_to_text(temp_file_path)
        else :

             return jsonify({"error": str(e)}, status_code=500)

        # Remove the temporary file
        os.remove(temp_file_path)

        return jsonify({"text": extracted_text})
    except Exception as e:
        return jsonify({"error": str(e)}, status_code=500)



@app.route('/drop', methods=['POST'])
def drop():
    request_data = request.get_json()

    api = request_data["api"]

    if not api.strip():  # This checks if the API key is empty or contains only spaces
    # Use the default API key
        default_api_key = "b60fe92d735dfe57981b989f910d19d14805b5cae8ab09b1e2fe052f8f41a437"  # Replace with your default API key
        configuration = Configuration(
        username=default_api_key
    )
    else:
        configuration = Configuration(
        username=api
    )

        # Initialize Dropbox API client
        with ApiClient(configuration) as api_client:
            signature_request_api = apis.SignatureRequestApi(api_client)
          # Parse JSON data from the request body
            text_data = request_data["chat"]
            pdf_file_path = text_to_docx(text_data)

            # Extract signer email addresses from the request data
            signer_1_email = request_data["signer_1_email"]
            signer_2_email = request_data["signer_2_email"]
            title = request_data["title"]
            subject = request_data["subject"]
            message = request_data["message"]
            cc_email_addresses = request_data["cc_email_addresses"]  # Retrieve cc_email_addresses as specified in the request data

        
            # Define signers and other options
            signer_1 = models.SubSignatureRequestSigner(
                email_address=signer_1_email,
                name=signer_1_email,
                order=0,
            )

            signer_2 = models.SubSignatureRequestSigner(
                email_address=signer_2_email,
                name=signer_2_email,
                order=1,
            )

            signing_options = models.SubSigningOptions(
                draw=True,
                type=True,
                upload=True,
                phone=True,
                default_type="draw",
            )

            field_options = models.SubFieldOptions(
                date_format="DD - MM - YYYY",
            )

            data = models.SignatureRequestSendRequest(
                title=title,
                subject=subject,
                message=message,
                signers=[signer_1, signer_2],
                cc_email_addresses=cc_email_addresses,
                files=[open(pdf_file_path, "rb")],  # Use the generated PDF
                metadata={
                    "custom_id": 1234,
                    "custom_text": "NDA #9",
                },
                signing_options=signing_options,
                field_options=field_options,
                test_mode=True,
            )
            try:

            # Send a signature request
                response = signature_request_api.signature_request_send(data)
                # print(response)

                return 'Check your inbox on your email for signing', 200

            except ApiException as e:
                print("Exception when calling Dropbox Sign API: %s\n" % e)
                return jsonify({'error': str(e)}, status_code=500)



@app.route('/update' , methods = ["POST"] )
def saveId():
    data = request.get_json()


    google_id = data["google_id"]
    conv_id = data["conv_id"]
    response = data["response"]

    data_to_upsert = {
        "googleid": google_id,
        "conv_id": conv_id,
        "response": response
    }
    try:
        # Attempt to upsert the data into the "Con" table
        supabase.from_("demo").upsert([data_to_upsert]).execute()
        return jsonify({"message": "Data upserted successfully"})
    except Exception as e:
        # Handle the exception and provide an appropriate error response
        return jsonify({"error": str(e)}), 500  # HTTP 500 Inter




@app.route('/get_conversations/<google_id>', methods=["GET"])
def getConversations(google_id):
    print(request.url)
    try:
        # Fetch data from the "demo" table based on the provided Google ID
        query = supabase.from_("demo").select("conv_id, response").eq("googleid", google_id)
        response = query.execute()

        # Check if the response contains data
        if response.data:
            rows = response.data
            # print(rows)
            # Create a dictionary to store conv_id as keys and lists of responses as values
            conv_id_responses = {}
            for row in rows:
                conv_id = row["conv_id"]
                response = row["response"]
                if conv_id not in conv_id_responses:
                    conv_id_responses[conv_id] = []
                conv_id_responses[conv_id].append(response)

            return jsonify(conv_id_responses)
    
        else:
            return jsonify({})  # No data found, return an empty dictionary
    except Exception as e:
        return jsonify({"error": str(e)}), 500  # HTTP 500 Internal Server Error for failure

@app.route('/chat', methods=['POST'])
def api():
    data = request.get_json()
    # input_message is the actual data, the data mime type is specified in type
    input_message = data['prompt']

 
    # ai_id is the id of the ai example GPT4 or GPT3.5 or LLAMA etc 
    conv_id= data["conversationId"]

   
    
    response = handle_message(input_message, conv_id,internet=False,assesment=False,spell=False)


    return jsonify({'response': response})

@app.route('/spell', methods=['POST'])
def spell():
    data = request.get_json()
    # input_message is the actual data, the data mime type is specified in type
    input_message = data['prompt']

 
    # ai_id is the id of the ai example GPT4 or GPT3.5 or LLAMA etc 
    conv_id= data["conversationId"]

   
    
    response = handle_message(input_message, conv_id,internet=False,spell=True,assesment=False)


    return jsonify({'response': response})

@app.route('/assessment', methods=['POST'])
def assessment():
    data = request.get_json()
    # input_message is the actual data, the data mime type is specified in type
    input_message = data['prompt']

 
    # ai_id is the id of the ai example GPT4 or GPT3.5 or LLAMA etc 
    conv_id= data["conversationId"]

   
    
    response = handle_message(input_message, conv_id,internet=False,assesment=True,spell=False)


    return jsonify({'response': response})

@app.route('/chat-internet', methods=['POST'])
def internet():
    data = request.get_json()
    # input_message is the actual data, the data mime type is specified in type
    input_message = data['prompt']

 
    # ai_id is the id of the ai example GPT4 or GPT3.5 or LLAMA etc 
    conv_id= data["conversationId"]

   
    
    response = handle_message(input_message, conv_id,internet=True,assesment=False,spell=False)


    return jsonify({'response': response})


@app.route('/delete-conv', methods=['DELETE'])
def deleteConversation():
    # Check if conversationId is provided
        
    data = request.get_json()
    conversationId = data['conversationId']
    if not conversationId:
        return jsonify({"message": "Invalid conversation ID provided"}, status_code=400)


    try:
        # Delete all rows in the "demo" table with the specified conversation ID
        supabase.from_("demo").delete().eq("conv_id", conversationId).execute()
        return jsonify({"message": f"Deleted all rows with conversation ID: {conversationId}"})
    except Exception as e:
        return jsonify({"message": f"Error deleting conversation: {str(e)}"}, status_code=500)






if __name__ == '__main__':
    app.run(host='0.0.0.0', port=os.environ.get('PORT', 4000), debug=False)
    
